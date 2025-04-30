import streamlit as st
import os
import tempfile
import openai
import datetime
from docx import Document
import zipfile
import io
import re
from io import BytesIO

openai.api_key = st.secrets["OPENAI_API_KEY"]

st.set_page_config(layout="wide")
st.title("📘 AI Training Content App (TTS verified)")

if st.sidebar.button("Test TTS"):
    test_file = os.path.join(tempfile.gettempdir(), "sidebar_test.mp3")
    response = openai.audio.speech.create(
        model="tts-1",
        voice="alloy",
        input="This is a sidebar TTS test inside your deployed app."
    )
    with open(test_file, "wb") as f:
        f.write(response.content)

    with open(test_file, "rb") as f:
        st.sidebar.download_button("Download TTS Test Audio", f, file_name="test.mp3")
        st.sidebar.audio(f.read(), format="audio/mp3")

uploaded_file = st.file_uploader("Upload a Word document", type=["docx"])

def extract_sections_from_docx(doc):
    sections = []
    current_heading = None
    current_text = []

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading 7"):
            if current_heading and current_text:
                sections.append((current_heading, "\n".join(current_text).strip()))
            current_heading = para.text.strip()
            current_text = []
        elif current_heading:
            current_text.append(para.text.strip())

    if current_heading and current_text:
        sections.append((current_heading, "\n".join(current_text).strip()))

    return sections

def create_word_doc(text, filename):
    doc = Document()
    doc.add_paragraph(text)
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    doc.save(temp_path)
    return temp_path

def create_text_file(text, filename):
    temp_path = os.path.join(tempfile.gettempdir(), filename)
    with open(temp_path, "w", encoding="utf-8") as f:
        f.write(text)
    return temp_path

def load_qref_template():
    return Document("templates/QREF_Template.docx")

def clear_document_after_table(doc):
    if doc.tables:
        last_table_element = doc.tables[0]._element
        following = list(last_table_element.itersiblings())
        for element in following:
            element.getparent().remove(element)

import subprocess

def create_audio_file(text, filename):
    output_path = os.path.join(tempfile.gettempdir(), filename)

    subprocess.run(
        ["python", "generate_tts.py", text, output_path],
        check=True
    )

    return output_path if os.path.exists(output_path) and os.path.getsize(output_path) > 0 else None



if uploaded_file:
    docx_stream = BytesIO(uploaded_file.read())
    doc = Document(docx_stream)
    extracted = extract_sections_from_docx(doc)

    if extracted:
        all_sections = [f"{i+1}. {title}" for i, (title, _) in enumerate(extracted)]
        selected_sections = st.multiselect("Choose section(s) to create class from", all_sections)

        col1, col2 = st.columns(2)
        with col1:
            if st.button("Create QuickByte"):
                st.session_state.run_type = "QuickByte"
        with col2:
            if st.button("Create FastTrack", disabled=len(selected_sections) < 1):
                st.session_state.run_type = "FastTrack"

        if "run_type" in st.session_state and st.session_state.run_type:
            selected_indices = [int(s.split(".")[0]) - 1 for s in selected_sections]
            selected_content = "\n\n".join(extracted[i][1] for i in selected_indices)
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

            with st.spinner("Generating content..."):
                outline_response = openai.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {"role": "system", "content": "You are an expert instructional designer."},
                        {"role": "user", "content": f"Create an outline with learning objectives for a {'15-minute QuickByte' if st.session_state.run_type == 'QuickByte' else '30-minute FastTrack'} instructor-led class based on this:\n{selected_content}"},
                    ]
                )

                script_response = openai.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {"role": "system", "content": "You are a professional e-learning narrator."},
                        {"role": "user", "content": f"Write a friendly but professional narration script for a video based on this content:\n{selected_content}"},
                    ]
                )

                tips_response = openai.chat.completions.create(
                    model="gpt-4.1-mini",
                    messages=[
                        {"role": "system", "content": "You are an instructional content expert."},
                        {"role": "user", "content": f"Generate exactly 5 email tips based on the following training content. Each tip should describe one useful feature, explain its benefit, and provide short step-by-step instructions. Output each as a separate tip, clearly numbered or titled:\n{selected_content}"},
                    ]
                )

                outline_text = outline_response.choices[0].message.content
                script_text = script_response.choices[0].message.content
                tips_text = tips_response.choices[0].message.content

                outline_file = create_word_doc(outline_text, f"class_outline_{timestamp}.docx")
                script_file = create_text_file(script_text, f"narration_script_{timestamp}.txt")

                paragraphs = script_text.split("\n\n")
                audio_files = []
                for idx, paragraph in enumerate(paragraphs):
                    audio_filename = f"narration_paragraph_{idx + 1}_{timestamp}.mp3"
                    audio_file = create_audio_file(paragraph, audio_filename)
                    audio_files.append(audio_file)

                tips = re.split(r"(?m)^\s*(?:\d+\.\s+|Tip\s+\d+:)", tips_text)
                tips = [tip.strip() for tip in tips if tip.strip()][:5]

                tip_files = []
                for i, tip in enumerate(tips):
                    tip_filename = f"email_tip_{i+1}_{timestamp}.docx"
                    tip_path = create_word_doc(tip, tip_filename)
                    tip_files.append((f"Email Tip {i+1}", tip, tip_path))

                tip_zip = io.BytesIO()
                with zipfile.ZipFile(tip_zip, "w") as zipf:
                    for _, _, path in tip_files:
                        zipf.write(path, os.path.basename(path))
                tip_zip.seek(0)

                qref_doc = load_qref_template()
                clear_document_after_table(qref_doc)

                class_title = selected_sections[0] if selected_sections else "Training Content"
                qref_doc.add_paragraph(class_title, style="IT Title")
                qref_doc.add_paragraph("Overview", style="IT Heading 1")
                qref_doc.add_paragraph("This Quick Reference supports the class learning objectives.", style="Body Text")

                for section in selected_content.split("\n\n"):
                    qref_doc.add_paragraph(section.strip(), style="IT Heading 2")
                    qref_doc.add_paragraph("[Insert Screenshot Here]", style="Body Text")

                for para in script_text.split("\n\n"):
                    if para.lower().startswith("tip:"):
                        qref_doc.add_paragraph(para.replace("TIP:", "").strip(), style="IT Tip")
                    elif para.lower().startswith("note:"):
                        qref_doc.add_paragraph(para.replace("NOTE:", "").strip(), style="IT Note")
                    else:
                        qref_doc.add_paragraph(para.strip(), style="IT Number_1")

                qref_filename = f"{class_title} QREF_{timestamp}.docx".replace("|", "-").replace(":", "-").replace("/", "-")
                qref_path = os.path.join(tempfile.gettempdir(), qref_filename)
                qref_doc.save(qref_path)

                st.session_state.generated = True
                st.session_state.timestamp = timestamp
                st.session_state.tabs = {
                    "Outline": (outline_text, outline_file),
                    "Narration": (script_text, script_file, audio_files),
                    "Email Tips": (tips, tip_zip),
                    "Quick Reference": (qref_path,)
                }

if st.session_state.get("generated"):
    tabs = st.tabs(["Outline", "Narration", "Email Tips", "Quick Reference"])

    with tabs[0]:
        tab_content, tab_file = st.session_state.tabs["Outline"]
        with open(tab_file, "rb") as f:
            st.download_button("Download Class Outline", data=f, file_name=os.path.basename(tab_file))
        st.markdown(tab_content)

    with tabs[1]:
        tab_content, tab_file, audio_files = st.session_state.tabs["Narration"]
        with open(tab_file, "rb") as f:
            st.download_button("Download Narration Script", data=f, file_name=os.path.basename(tab_file))
        st.markdown(tab_content)
        for audio_file in audio_files:
            if audio_file and os.path.exists(audio_file):
                with open(audio_file, "rb") as af:
                    st.download_button(f"Download Narration Audio (mp3)", data=af, file_name=os.path.basename(audio_file))
                    st.audio(af.read(), format="audio/mp3")

    with tabs[2]:
        tip_texts, tip_zip = st.session_state.tabs["Email Tips"]
        st.download_button("Download All Email Tips", data=tip_zip, file_name=f"email_tips_{st.session_state.timestamp}.zip")
        for i, tip in enumerate(tip_texts):
            st.markdown(f"**Tip {i+1}:** {tip}")

    with tabs[3]:
        qref_path, = st.session_state.tabs["Quick Reference"]
        with open(qref_path, "rb") as f:
            st.download_button("Download Quick Reference", data=f, file_name=os.path.basename(qref_path))
