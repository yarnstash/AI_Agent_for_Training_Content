import streamlit as st
import os
import tempfile
import openai
import datetime
from docx import Document
from docx.shared import Inches
from io import BytesIO

st.set_page_config(layout="wide")
st.title("AI Training Content App QREF")

openai.api_key = st.secrets["OPENAI_API_KEY"]

uploaded_files = st.file_uploader("Upload one or more source documents (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

document_chunks = []

def extract_text_from_pdf(file_path):
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

if uploaded_files:
    for uploaded_file in uploaded_files:
        suffix = ".pdf" if uploaded_file.name.endswith(".pdf") else ".docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name

        if suffix == ".pdf":
            extracted_text = extract_text_from_pdf(tmp_path)
        else:
            extracted_text = extract_text_from_docx(tmp_path)

        document_chunks.append((uploaded_file.name, extracted_text))

    st.success("Files uploaded and content extracted.")

    query = st.text_input("What content are you looking for?")
    if query and document_chunks and "search_topics" not in st.session_state:
        with st.spinner("Finding relevant topics..."):
            combined_text = "  ".join([text for _, text in document_chunks])
            user_prompt = f"""From this content:

{combined_text}

What topics are relevant to this query: {query}
"""
            response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are a document analyst. Extract a list of specific, self-contained topics based on the query. Format as a numbered or bullet list."},
                    {"role": "user", "content": user_prompt}
                ]
            )
            topics = response.choices[0].message.content
            topic_lines = [line.strip("â€¢-1234567890. ") for line in topics.strip().splitlines() if line.strip()]
            st.session_state.search_topics = topic_lines
            st.session_state.full_text = combined_text

if "search_topics" in st.session_state:
    st.markdown("### Step 2: Choose Topics for Your Class")
    selected = st.multiselect("Pick the topics you'd like to include:", st.session_state.search_topics)
    if selected and "selected_text" not in st.session_state:
        with st.spinner("Extracting selected content..."):
            selection_prompt = f"""Extract detailed content for these selected topics:

{selected}

From the following documents:

{st.session_state.full_text}
"""
            content_response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You extract training content from documents."},
                    {"role": "user", "content": selection_prompt}
                ]
            )
            st.session_state.selected_text = content_response.choices[0].message.content
            st.success("Content extracted.")

    if selected and "selected_text" in st.session_state:
        if st.button("Generate QREF"):
            full_text = st.session_state.selected_text.strip()
            lines = full_text.splitlines()
            overview_lines, body_lines = [], []
            for i, line in enumerate(lines):
                if line.strip().startswith("### "):
                    overview_lines = lines[:i]
                    body_lines = lines[i:]
                    break
            else:
                overview_lines = lines[:2]
                body_lines = lines[2:]

            overview = " ".join(overview_lines).strip()
            steps = "\n".join(body_lines).strip()
            selected_topic = ", ".join(selected)
            today = datetime.date.today().strftime("%B %d, %Y")

            def safe_style(doc, style_name, fallback="Normal"):
                try:
                    doc.styles[style_name]
                    return style_name
                except KeyError:
                    return fallback

            def clear_below_first_table(doc):
                first_table = doc.tables[0]
                last_tbl_elm = first_table._element
                body_elm = doc._body._element
                to_delete = []
                found_table = False
                for child in body_elm.iterchildren():
                    if child == last_tbl_elm:
                        found_table = True
                        continue
                    if found_table:
                        to_delete.append(child)
                for element in to_delete:
                    body_elm.remove(element)

            def create_qref_docx(app, function, audience, version, overview, steps, tips, related, template_path):
                doc = Document(template_path)
                clear_below_first_table(doc)

                # Set margins safely if available
                if doc.sections:
                    section = doc.sections[0]
                    section.top_margin = Inches(0.5)
                    section.bottom_margin = Inches(0.6)
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                doc.add_paragraph("OVERVIEW", style=safe_style(doc, "IT Heading 1"))
                doc.add_paragraph(overview, style=safe_style(doc, "IT Body Text"))

                current_list = None
                for line in steps.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    if line.startswith("### "):
                        doc.add_paragraph(line.replace("###", "").strip(), style=safe_style(doc, "IT Heading 1"))
                        current_list = None  # reset list
                    elif line.startswith("- "):
                        current_list = doc.add_paragraph(line.strip("- ").strip(), style=safe_style(doc, "IT Number_1"))
                    else:
                        current_list = doc.add_paragraph(line, style=safe_style(doc, "IT Number_1"))

                doc.add_paragraph("TIPS & NOTES", style=safe_style(doc, "IT Heading 1"))
                for tip in tips:
                    doc.add_paragraph(tip, style=safe_style(doc, "IT Tip"))

                doc.add_paragraph("RELATED FEATURES", style=safe_style(doc, "IT Heading 1"))
                for item in related:
                    doc.add_paragraph(item, style=safe_style(doc, "IT Note"))

                output = BytesIO()
                doc.save(output)
                output.seek(0)
                return output

            template_path = os.path.join("templates", "QREF_Template.docx")
            docx_file = create_qref_docx(
                app="[App Name]",
                function="[Function or Module]",
                audience="[User Type]",
                version=today,
                overview=overview,
                steps=steps,
                tips=["[Add any helpful tips or reminders.]"],
                related=["[Mention other relevant QREFs or tools.]"],
                template_path=template_path
            )

            st.download_button(
                label="Download QREF Word Document",
                data=docx_file,
                file_name=f"QREF_{selected_topic}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )