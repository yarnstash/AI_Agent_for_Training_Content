import streamlit as st
import os
import tempfile
import openai
import datetime
from docx import Document
from PyPDF2 import PdfReader
import zipfile
import io
import re

openai.api_key = st.secrets["OPENAI_API_KEY"]

st.set_page_config(layout="wide")
st.title("ðŸ“˜ AI Training Content App with Search + Topic Picker")

uploaded_files = st.file_uploader("Upload one or more source documents (PDF or DOCX)", type=["pdf", "docx"], accept_multiple_files=True)

selected_sections = []
all_sections = []
document_chunks = []

def extract_text_from_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    return " ".join(p.text for p in doc.paragraphs if p.text.strip())

if uploaded_files:
    for uploaded_file in uploaded_files:
        suffix = ".pdf" if uploaded_file.name.endswith(".pdf") else ".docx"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
            tmp_file.write(uploaded_file.read())
            tmp_path = tmp_file.name

        if suffix == ".pdf":
            extracted_text = extract_text_from_pdf(tmp_path)
        else:
            extracted_text = extract_text_from_docx(tmp_path)

        document_chunks.append((uploaded_file.name, extracted_text))

    st.success("Files uploaded and content extracted.")

    query = st.text_input("What content are you looking for?")
    if query and document_chunks and "search_topics" not in st.session_state:
        with st.spinner("Finding relevant topics..."):
            combined_text = "  ".join([text for _, text in document_chunks])
            user_prompt = f"""From this content:

{combined_text}

What topics are relevant to this query: {query}
"""
            response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are a document analyst. Extract a list of specific, self-contained topics based on the query. Format as a numbered or bullet list."},
                    {"role": "user", "content": user_prompt}
                ]
            )
            topics = response.choices[0].message.content
            topic_lines = [line.strip("â€¢-1234567890. ") for line in topics.strip().splitlines() if line.strip()]
            st.session_state.search_topics = topic_lines
            st.session_state.full_text = combined_text

if "search_topics" in st.session_state:
    st.markdown("### Step 2: Choose Topics for Your Class")
    selected = st.multiselect("Pick the topics you'd like to include:", st.session_state.search_topics)
    if selected:
        with st.spinner("Extracting selected content..."):
            selection_prompt = f"""Extract detailed content for these selected topics:

{selected}

From the following documents:

{st.session_state.full_text}
"""
            content_response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You extract training content from documents."},
                    {"role": "user", "content": selection_prompt}
                ]
            )
            st.session_state.selected_text = content_response.choices[0].message.content
            st.success("Content extracted.")

if "selected_text" in st.session_state:
    st.markdown("### Step 3: Generate Training Materials")
    selected_text = st.session_state.selected_text
    st.text_area("Selected Content", selected_text, height=200)

    col1, col2 = st.columns(2)
    if col1.button("Create QuickByte"):
        st.session_state.run_type = "QuickByte"
    if col2.button("Create FastTrack"):
        st.session_state.run_type = "FastTrack"

    if "run_type" in st.session_state:
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        with st.spinner("Generating training content..."):
            outline_prompt = f"""Create a detailed outline with learning objectives for a {st.session_state.run_type} class based on this:

{selected_text}
"""
            script_prompt = f"""Write a narration script for a video class based on this:

{selected_text}
"""
            tips_prompt = f"""Generate 5 email tips based on this content. Each tip should be clearly separated by 'Tip X:' and include a benefit and step-by-step instructions:

{selected_text}
"""
            outline_response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are an expert instructional designer."},
                    {"role": "user", "content": outline_prompt}
                ]
            )
            script_response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are a professional training narrator."},
                    {"role": "user", "content": script_prompt}
                ]
            )
            tips_response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are an expert trainer."},
                    {"role": "user", "content": tips_prompt}
                ]
            )

            outline = outline_response.choices[0].message.content
            script = script_response.choices[0].message.content
            tips_text = tips_response.choices[0].message.content

            def save_docx(text, filename):
                path = os.path.join(tempfile.gettempdir(), filename)
                doc = Document()
                doc.add_paragraph(text)
                doc.save(path)
                return path

            def save_txt(text, filename):
                path = os.path.join(tempfile.gettempdir(), filename)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(text)
                return path

            outline_file = save_docx(outline, f"outline_{timestamp}.docx")
            script_file = save_txt(script, f"script_{timestamp}.txt")

            tip_blocks = re.findall(r"Tip\s+\d+:(.*?)(?=Tip\s+\d+:|\Z)", tips_text, re.DOTALL)
            tips = [f"Tip {i+1}: {block.strip()}" for i, block in enumerate(tip_blocks)][:5]

            tip_paths = []
            for i, tip in enumerate(tips):
                tip_paths.append(save_docx(tip, f"email_tip_{i+1}_{timestamp}.docx"))

            st.session_state.tabs = {
                "Outline": (outline, outline_file),
                "Narration": (script, script_file),
                "Email Tips": (tips, tip_paths)
            }

if "tabs" in st.session_state:
    tabs = st.tabs(["Outline", "Narration", "Email Tips"])

    with tabs[0]:
        content, path = st.session_state.tabs["Outline"]
        with open(path, "rb") as f:
            st.download_button("Download Outline", f, os.path.basename(path))
        st.markdown(content)

    with tabs[1]:
        content, path = st.session_state.tabs["Narration"]
        with open(path, "rb") as f:
            st.download_button("Download Narration Script", f, os.path.basename(path))
        st.markdown(content)

    with tabs[2]:
        tips, paths = st.session_state.tabs["Email Tips"]
        for i, (tip, path) in enumerate(zip(tips, paths)):
            st.markdown(f"**Tip {i+1}:** {tip}")
            with open(path, "rb") as f:
                st.download_button(f"Download Tip {i+1}", f, os.path.basename(path))