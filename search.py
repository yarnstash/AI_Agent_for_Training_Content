import streamlit as st
import os
import tempfile
from docx import Document
from PyPDF2 import PdfReader
import openai
from io import BytesIO

# Set your API key from Streamlit secrets
openai.api_key = st.secrets["OPENAI_API_KEY"]

# App layout
st.set_page_config(layout="wide")
st.title("Semantic Document Search")

# File uploader
uploaded_files = st.file_uploader("Upload PDF or Word documents", type=["pdf", "docx"], accept_multiple_files=True)

documents = []

# Text extraction
def extract_text_from_pdf(file):
    reader = PdfReader(file)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join(para.text for para in doc.paragraphs)

# Process uploaded files
if uploaded_files:
    for file in uploaded_files:
        if file.name.endswith(".pdf"):
            text = extract_text_from_pdf(file)
        elif file.name.endswith(".docx"):
            text = extract_text_from_docx(file)
        else:
            continue
        documents.append({"filename": file.name, "content": text})
    st.success(f"Loaded {len(documents)} documents.")

# Text input
user_query = st.text_input("What would you like to extract or search for?")

if user_query and documents:
    full_text = "\n\n".join(doc["content"] for doc in documents)

    with st.spinner("Querying documents..."):
        try:
            response = openai.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant who extracts relevant content from documents."},
                    {"role": "user", "content": f"Search the following document content and respond to this prompt:\n\n{user_query}\n\nDOCUMENTS:\n{full_text}"}
                ]
            )
            result = response.choices[0].message.content
        except Exception as e:
            st.error(f"OpenAI API error: {e}")
            result = ""

    if result:
        st.subheader("Search Result")
        st.markdown(result)

        # Save to Word
        output_path = os.path.join(tempfile.gettempdir(), "search_result.docx")
        doc = Document()
        for line in result.split("\n"):
            doc.add_paragraph(line.strip())
        doc.save(output_path)

        with open(output_path, "rb") as f:
            st.download_button("Download as Word Document", f, file_name="search_result.docx")